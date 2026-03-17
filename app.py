import streamlit as st
import os
from crewai import Agent, Task, Crew, Process, LLM
import docx
import io

# 1. 웹사이트 기본 설정
st.set_page_config(page_title="AI 전략 기획실", page_icon="🚀", layout="wide")
st.title("🚀 대상주식회사 AI 전략 기획실 (Red Team)")
st.markdown("비즈니스 주제와 특별 요청사항을 입력하면 4명의 AI 전문가가 치열하게 토론하여 최종 워드 보고서를 작성합니다.")

# 2. 사이드바: 보안을 위한 API 키 입력창
with st.sidebar:
    st.header("🔑 API 키 설정")
    google_api_key = st.text_input("Google Gemini API Key", type="password")
    groq_api_key = st.text_input("Groq API Key (Llama 3.3)", type="password")
    st.markdown("---")
    st.markdown("이 시스템은 보안을 위해 API 키를 저장하지 않습니다. 실행 시에만 사용됩니다.")

# 3. 메인 화면: 주제 입력
topic = st.text_input("▶ 논의할 비즈니스 주제를 입력하세요", placeholder="예: 간편식 신제품 편의점 입점 전략")
requirement = st.text_input("▶ 반드시 포함해야 할 특별 요청사항", placeholder="예: 마케팅 예산 최소화, 2030 타겟")

# 4. 실행 버튼 및 AI 엔진 가동
if st.button("AI 토론 시작 및 보고서 생성", type="primary"):
    if not google_api_key or not groq_api_key:
        st.error("좌측 사이드바에 두 개의 API 키를 모두 입력해 주세요.")
    elif not topic:
        st.warning("비즈니스 주제를 입력해 주세요.")
    else:
        # 스피너(로딩 애니메이션) 표시
        with st.spinner('AI 드림팀이 시장을 분석하고 치열하게 토론 중입니다. (약 1~3분 소요)...'):
            os.environ["GOOGLE_API_KEY"] = google_api_key
            os.environ["GROQ_API_KEY"] = groq_api_key

            # LLM 및 에이전트 세팅
            gemini_llm = LLM(model="gemini/gemini-1.5-flash", temperature=0.7)
            groq_llm = LLM(model="groq/llama-3.3-70b-versatile", api_key=groq_api_key, temperature=0.3)

            researcher = Agent(role='시장 데이터 분석가', goal='최신 정량적 데이터 수집', backstory='팩트와 수치로만 말하는 차가운 전문가', llm=groq_llm)
            strategist = Agent(role='사업 기획실장', goal='경쟁사를 압도할 파격적인 기획안 작성', backstory='공격적인 전략을 밀어붙이는 기획자', llm=gemini_llm)
            critic = Agent(role='리스크 수문장', goal='결함 3가지를 반드시 찾아내고 공격', backstory='기획안을 비판하는 재무/법무 책임자', llm=groq_llm)
            editor = Agent(role='의사결정권자', goal='최적의 타협안을 한글 보고서로 작성', backstory='냉철한 비즈니스 문서 조정자', llm=gemini_llm)

            # 태스크 세팅
            t1 = Task(description=f"[{topic}] 시장 조사. 특별 요청사항: [{requirement}]", agent=researcher, expected_output='데이터 요약본')
            t2 = Task(description=f"[{topic}] 영업/마케팅 플랜. 특별 요청사항: [{requirement}]", agent=strategist, context=[t1], expected_output='실행 전략')
            t3 = Task(description=f"[{topic}] 전략에 대한 치명적 결함 비판.", agent=critic, context=[t2], expected_output='비판 보고서')
            t4 = Task(description=f"[{topic}] 최종 실행 보고서 작성.", agent=editor, context=[t2, t3], expected_output='최종 제안서')

            # 토론 실행
            crew = Crew(agents=[researcher, strategist, critic, editor], tasks=[t1, t2, t3, t4], process=Process.sequential)
            result = crew.kickoff()
            final_text = result.raw

            st.success("토론 및 최종 보고서 작성이 완료되었습니다!")
            
            # 결과 화면 출력
            st.markdown("### 🎯 [최종 비즈니스 제안서]")
            st.markdown(final_text)

            # 워드 파일 변환 및 다운로드 버튼 생성
            doc = docx.Document()
            doc.add_heading(f'최종 비즈니스 제안서: {topic}', 0)
            doc.add_paragraph(final_text)
            
            bio = io.BytesIO()
            doc.save(bio)
            
            st.download_button(
                label="📄 워드 파일로 다운로드",
                data=bio.getvalue(),
                file_name="대상주식회사_전략보고서.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
